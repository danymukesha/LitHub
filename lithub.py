import base64
import os
import re
import sqlite3
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from flask import Flask, flash, redirect, render_template_string, request, url_for
from PIL import Image

app = Flask(__name__)
app.secret_key = "your_secret_key"

REVIEWS_DIR = "reviews"
DB_FILE = "comments.db"
os.makedirs(REVIEWS_DIR, exist_ok=True)

def init_db():
    with sqlite3.connect(DB_FILE) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS comments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                review_name TEXT,
                comment TEXT,
                timestamp TEXT
            )
        """)

init_db()

def extract_docx_content(filepath):
    def iter_block_items(parent):
        for child in parent.element.body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield Table(child, parent)

    doc = Document(filepath)
    title = doc.paragraphs[0].text.strip() if doc.paragraphs else "Untitled Review"
    description = doc.paragraphs[1].text.strip() if len(doc.paragraphs) > 1 else "No description available."
    blocks = list(iter_block_items(doc))[2:]

    content = []
    list_open = None

    for block in blocks:
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name.lower()

            if "heading" in style or text.lower().startswith("chapter"):
                if list_open:
                    content.append(f"</{list_open}>")
                    list_open = None
                content.append(f'<h2 class="text-xl font-bold mt-6 mb-2">{text}</h2>')
                continue

            if style.startswith("list") or re.search(r"bullet|number", style):
                tag = "ul" if "bullet" in style else "ol"
                if list_open != tag:
                    if list_open:
                        content.append(f"</{list_open}>")
                    content.append(f"<{tag}>")
                    list_open = tag
                content.append(f"<li>{text}</li>")
                continue

            if list_open:
                content.append(f"</{list_open}>")
                list_open = None

            para_html = ""
            for run in block.runs:
                run_text = run.text
                if run_text:
                    if run.bold:
                        run_text = f"<b>{run_text}</b>"
                    if run.italic:
                        run_text = f"<i>{run_text}</i>"
                    para_html += run_text

            if para_html.strip():
                content.append(f"<p>{para_html.strip()}</p>")

        elif isinstance(block, Table):
            if list_open:
                content.append(f"</{list_open}>")
                list_open = None
            table_html = '<table class="table-auto border border-collapse border-gray-300 my-4">'
            for row in block.rows:
                table_html += "<tr>" + "".join(
                    f'<td class="border p-2">{cell.text.strip()}</td>' for cell in row.cells
                ) + "</tr>"
            table_html += "</table>"
            content.append(table_html)

    if list_open:
        content.append(f"</{list_open}>")

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            img = Image.open(BytesIO(img_data))
            fmt = (img.format or "JPEG").lower()
            img_io = BytesIO()
            img.save(img_io, format=img.format)
            img_base64 = base64.b64encode(img_io.getvalue()).decode()
            content.append(f'<img src="data:image/{fmt};base64,{img_base64}" class="my-4 max-w-full h-auto"/>')

    return title, description, "".join(content), []

@app.route("/", methods=["GET", "POST"])
def home():
    query = request.args.get("query", "").lower()
    reviews = [f for f in os.listdir(REVIEWS_DIR) if f.endswith(".docx")]

    reviews_data = []
    for review in reviews:
        path = os.path.join(REVIEWS_DIR, review)
        title, description, _, _ = extract_docx_content(path)
        if query in title.lower() or query in description.lower():
            created_on = datetime.fromtimestamp(os.path.getctime(path)).strftime("%Y-%m-%d %H:%M:%S")
            reviews_data.append({
                "filename": review,
                "title": title,
                "description": description,
                "created_on": created_on,
            })

    if request.method == "POST":
        file = request.files.get("file")
        if file and file.filename.lower().endswith(".docx"):
            file.save(os.path.join(REVIEWS_DIR, file.filename))
            flash("File uploaded successfully!", "success")
            return redirect(url_for("home"))
        flash("Invalid file format. Please upload a .docx file.", "danger")

    no_results = not reviews_data
    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Systematic Literature Reviews</title>
        <script src="https://cdn.tailwindcss.com"></script>
    </head>
    <body class="bg-gray-100 font-sans">
        <header class="bg-blue-600 text-white py-6">
            <div class="container mx-auto px-4">
                <h1 class="text-3xl font-bold">Systematic Literature Reviews</h1>
                <p class="mt-2">A collection of peer-reviewed systematic literature reviews</p>
            </div>
        </header>
        <main class="container mx-auto px-4 py-8">
            <div class="mb-6">
                <!-- Search Bar -->
                <form method="GET" action="/" class="flex items-center space-x-4">
                    <input type="text" name="query" value="{{ query }}" placeholder="Search reviews..."
                        class="w-full p-2 border rounded-lg" />
                    <button type="submit" class="bg-blue-600 text-white py-2 px-4 rounded-lg hover:bg-blue-700">
                        Search
                    </button>
                </form>
            </div>

            <div class="mb-6">
                {% if no_results %}
                    <div class="bg-yellow-100 text-yellow-800 p-3 rounded-lg">
                        No reviews found for your query.
                    </div>
                {% endif %}
            </div>

            <div class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6">
                {% for review in reviews %}
                <div class="bg-white p-6 rounded-lg shadow-md">
                    <h2 class="text-xl font-semibold mb-2">{{ review.title }}</h2>
                    <p class="text-gray-600 mb-4">{{ review.description }}</p>
                    <a href="/review/{{ review.filename }}" class="text-blue-600 hover:underline">
                    Read More
                    </a>
                    <p class="text-gray-500 text-sm">Created on {{ review.created_on }}</p> <!-- Add this line -->
                </div>
                {% endfor %}
            </div>

            <!-- File Upload Form -->
            <div class="mt-8">
                <h2 class="text-xl font-semibold mb-4">Upload New Review</h2>
                <form method="POST" enctype="multipart/form-data">
                    <input type="file" name="file" class="border p-2 rounded-lg mb-4" accept=".docx" required>
                    <button type="submit" class="bg-blue-600 text-white py-2 px-4 rounded-lg hover:bg-blue-700">
                        Upload Review
                    </button>
                </form>
                {% with messages = get_flashed_messages(with_categories=true) %}
                    {% if messages %}
                    <div class="mt-4">
                        {% for category, message in messages %}
                        <div class="bg-{{ category }}-100 text-{{ category }}-800 p-3 rounded-lg mb-3">
                            {{ message }}
                        </div>
                        {% endfor %}
                    </div>
                    {% endif %}
                {% endwith %}
            </div>
        </main>
        <footer class="bg-gray-800 text-white py-4">
            <div class="container mx-auto px-4 text-center">
                <p>&copy; 2025 Systematic Reviews Blog. All rights reserved.</p>
            </div>
        </footer>
    </body>
    </html>
    """ 
    return render_template_string(html, reviews=reviews_data, query=query, no_results=no_results)

@app.route("/review/<name>", methods=["GET", "POST"])
def review(name):
    if not name.lower().endswith(".docx"):
        return "Invalid file", 400
    path = os.path.join(REVIEWS_DIR, name)
    if not os.path.exists(path):
        return "File not found", 404

    if request.method == "POST":
        comment = request.form.get("comment")
        if comment:
            with sqlite3.connect(DB_FILE) as conn:
                conn.execute(
                    "INSERT INTO comments (review_name, comment, timestamp) VALUES (?, ?, ?)",
                    (name, comment, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                )
        return redirect(url_for("review", name=name))

    title, description, content, images = extract_docx_content(path)
    with sqlite3.connect(DB_FILE) as conn:
        comments = conn.execute(
            "SELECT comment, timestamp FROM comments WHERE review_name = ? ORDER BY timestamp DESC", (name,)
        ).fetchall()

    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{{ title }}</title>
        <script src="https://cdn.tailwindcss.com"></script>
    </head>
    <body class="bg-gray-100 font-sans">
        <header class="bg-blue-600 text-white py-6">
            <div class="container mx-auto px-4">
                <h1 class="text-3xl font-bold">{{ title }}</h1>
                <p class="mt-2">{{ description }}</p>
            </div>
        </header>
        <main class="container mx-auto px-4 py-8">
            <div class="bg-white p-6 rounded-lg shadow-md">
                <div class="prose max-w-none">
                  {{ content | safe }}
                  {% for img in images %}
                    <img src="{{ img }}" alt="Document Image" class="my-4 max-w-1/4 h-auto" style="width: 100px; height: auto;">
                  {% endfor %}
                  <a href="/edit/{{ name }}" class="text-green-600 hover:underline mt-4 inline-block
                    bg-green-600 text-white py-2 px-4 rounded-lg hover:bg-green-700 mt-4 inline-block">
                    Edit Review
                  </a>
                </div>
                <a href="/" class="text-blue-600 hover:underline mt-4 inline-block">Back to Home</a>
              </div>
            <div class="mt-8">
                <h2 class="text-2xl font-semibold mb-4">Comments</h2>
                <form method="POST" class="mb-6">
                    <textarea name="comment" rows="4" class="w-full p-3 border rounded-lg"
                        placeholder="Add your comment..." required></textarea>
                    <button type="submit" class="bg-blue-600 text-white py-2 px-4 rounded-lg
                        hover:bg-blue-700">Submit Comment</button>
                </form>
                {% if comments %}
                <div class="space-y-4">
                    {% for comment in comments %}
                    <div class="bg-gray-50 p-4 rounded-lg">
                        <p class="text-gray-800">{{ comment[0] }}</p>
                        <p class="text-gray-500 text-sm">{{ comment[1] }}</p>
                    </div>
                    {% endfor %}
                </div>
                {% else %}
                <p class="text-gray-600">No comments yet. Be the first to comment!</p>
                {% endif %}
            </div>
        </main>
        <footer class="bg-gray-800 text-white py-4">
            <div class="container mx-auto px-4 text-center">
                <p>&copy; 2025 Systematic Reviews Blog. All rights reserved.</p>
            </div>
        </footer>
    </body>
    </html>
    """
    return render_template_string(html, title=title, description=description, content=content, images=images, comments=comments, name=name)

@app.route("/edit/<filename>", methods=["GET", "POST"])
def edit_review(filename):
    path = os.path.join(REVIEWS_DIR, filename)
    if not os.path.exists(path):
        return "File not found", 404

    title, description, _, _ = extract_docx_content(path)

    if request.method == "POST":
        new_title = request.form.get("title")
        new_description = request.form.get("description")
        file = request.files.get("file")
        updated = False

        if new_title and new_description:
            try:
                doc = Document(path)
                while len(doc.paragraphs) < 2:
                    doc.add_paragraph("")
                doc.paragraphs[0].text = new_title
                doc.paragraphs[1].text = new_description
                doc.save(path)
                updated = True
            except Exception as e:
                flash(f"Failed to update title/description: {e}", "danger")
                return redirect(url_for("edit_review", filename=filename))

        if file and file.filename.lower().endswith(".docx"):
            try:
                file.save(path)
                updated = True
            except Exception as e:
                flash(f"Failed to save new file: {e}", "danger")
                return redirect(url_for("edit_review", filename=filename))

        if updated:
            flash("Review updated successfully!", "success")
            return redirect(url_for("review", name=filename))
        else:
            flash("No changes made. Update title/description or upload a new .docx file.", "warning")
            return redirect(url_for("edit_review", filename=filename))

    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8" />
        <meta name="viewport" content="width=device-width, initial-scale=1" />
        <title>Edit Review</title>
        <script src="https://cdn.tailwindcss.com"></script>
    </head>
    <body class="bg-gray-100 font-sans">
        <header class="bg-blue-600 text-white py-6">
            <div class="container mx-auto px-4">
                <h1 class="text-3xl font-bold">Edit Review: {{ title }}</h1>
            </div>
        </header>
        <main class="container mx-auto px-4 py-8">
            {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
                <div class="mb-4">
                    {% for category, message in messages %}
                        <div class="bg-{{ 'green' if category == 'success' else 'red' if category == 'danger' else 'yellow' }}-100 text-{{ 'green' if category == 'success' else 'red' if category == 'danger' else 'yellow' }}-800 p-3 rounded-lg mb-2">
                            {{ message }}
                        </div>
                    {% endfor %}
                </div>
            {% endif %}
            {% endwith %}
            <form method="POST" enctype="multipart/form-data">
                <div class="mb-4">
                    <label for="title" class="block text-gray-700">Title</label>
                    <input type="text" id="title" name="title" value="{{ title }}" required class="w-full p-2 border rounded-lg" />
                </div>
                <div class="mb-4">
                    <label for="description" class="block text-gray-700">Description</label>
                    <textarea id="description" name="description" rows="4" required class="w-full p-2 border rounded-lg">{{ description }}</textarea>
                </div>
                <p class="mb-4 text-gray-700">
                    You can also upload a new .docx file to replace the entire review content.
                </p>
                <div class="mb-4">
                    <label for="file" class="block text-gray-700">Upload New DOCX (Optional)</label>
                    <input type="file" id="file" name="file" accept=".docx" class="w-full p-2 border rounded-lg" />
                </div>
                <button type="submit" class="bg-blue-600 text-white py-2 px-4 rounded-lg hover:bg-blue-700">Save Changes</button>
            </form>
            <a href="{{ url_for('review', name=filename) }}" class="text-blue-600 hover:underline mt-4 inline-block">Back to Review</a>
        </main>
    </body>
    </html>
    """
    return render_template_string(html, title=title, description=description, filename=filename)

if __name__ == "__main__":
    app.run(debug=True)
